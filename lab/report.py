from calculations import Wallpaper, Tile, Laminate
from docx import Document
from openpyxl import Workbook

def generate_report(material, area, format):
    if material == "Обои":
        material_obj = Wallpaper(area)
    elif material == "Плитка":
        material_obj = Tile(area)
    elif material == "Ламинат":
        material_obj = Laminate(area)

    quantity = material_obj.calculate_quantity()
    cost = material_obj.calculate_cost()

    if format == "doc":
        document = Document()
        document.add_paragraph(f"Материал: {material}")
        document.add_paragraph(f"Площадь: {area}")
        document.add_paragraph(f"Количество: {quantity}")
        document.add_paragraph(f"Стоимость: {cost}")
        document.save("report.docx")
    elif format == "xls":
        workbook = Workbook()
        sheet = workbook.active
        sheet["A1"] = "Материал"
        sheet["B1"] = "Площадь"
        sheet["C1"] = "Количество"
        sheet["D1"] = "Стоимость"
        sheet.append([material, area, quantity, cost])
        workbook.save("report.xlsx")
